import os
import requests
import json
from docx import Document
import tempfile
import subprocess
from dotenv import load_dotenv

# Load API key
load_dotenv()
api_key = os.getenv("OPEN_ROUTER_API_KEY")
if not api_key:
    raise ValueError("OPEN_ROUTER_API_KEY not found.")

# Ask user
prompt = input("How can I help you today?> ")

# Call OpenRouter
url = "https://openrouter.ai/api/v1/chat/completions"
headers = {
    "Authorization": f"Bearer {api_key}",
    "Content-Type": "application/json",
    "HTTP-Referer": "http://localhost"
}
data = {
    "model": "mistralai/mistral-7b-instruct",
    "messages": [{"role": "user", "content": prompt}],
    "temperature": 0.7
}

response = requests.post(url, headers=headers, data=json.dumps(data))
response.raise_for_status()
result = response.json()
generated_text = result['choices'][0]['message']['content']

# Create a temp file path manually to avoid locking issue
temp_path = os.path.join(tempfile.gettempdir(), "preview.docx")
doc = Document()
doc.add_heading("Generated Content", level=1)
doc.add_paragraph(generated_text)
doc.save(temp_path)

# Open in MS Word
os.startfile(temp_path)

# Ask for confirmation
input("\nPreview opened in Word. Press ENTER to save to Desktop, or Ctrl+C to cancel...")

# Save to Desktop
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop", "generated_output.docx")
doc.save(desktop_path)
print(f"Saved to: {desktop_path}")
