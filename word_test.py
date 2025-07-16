from docx import Document
from docx2pdf import convert
import os

# Set output path
output_dir = "C:\\Temp"
os.makedirs(output_dir, exist_ok=True)

output_path = os.path.join(output_dir, "com_test.docx")

# Create .docx file
doc = Document()
doc.add_heading("Leave Letter", level=1)
doc.add_paragraph("Dear Sir/Madam,")
doc.add_paragraph("I am writing to request a leave of absence for 2 days due to personal reasons.")
doc.add_paragraph("Thank you for your consideration.")
doc.add_paragraph("Sincerely,\nSidd")
doc.save(output_path)

print(f".docx saved at: {output_path}")

# Convert to PDF
try:
    convert(output_path)
    print("PDF successfully created.")
except Exception as e:
    print("PDF conversion failed:")
    print(e)
